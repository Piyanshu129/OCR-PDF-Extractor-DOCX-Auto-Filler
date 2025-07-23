# import streamlit as st
# import pytesseract
# from pdf2image import convert_from_bytes
# from PIL import Image
# import os
# import re
# from datetime import datetime
# from docx import Document
# import zipfile

# users = {
#     "admin": "Admin2123",
#     "piyanshu": "PiY@2025_Secure!",
#     "test":"test"
# }

# if "user" not in st.session_state:
#     st.title("🔐 Login")
#     username = st.text_input("Username")
#     password = st.text_input("Password", type="password")

#     if st.button("Login"):
#         if username in users and users[username] == password:
#             st.session_state["user"] = username
#             st.rerun()
#         else:
#             st.error("Invalid credentials")
#     st.stop()















# # def extract_booking_number(text):
# # pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # # ------------------ Main App ------------------


# # st.title("📄 OCR PDF Extractor & DOCX Auto-Filler")

# # # Poppler path (for macOS)
# # poppler_path = r"C:\Program Files\poppler-24.08.0\Library\bin"
# # os.environ["PATH"] += os.pathsep + poppler_path

# #     """
# #     Extract booking numbers from various shipping line booking confirmations.
# #     Supports multiple formats from different carriers like KMTC, ANL, PIL, CMA CGM, etc.
# #     """
# #     patterns = [
# #         r"Portal\s*Booking\s*Ref\s*[:\-]?\s*(.*)",

# #                 # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
# #                 # E.g., "Booking Number : DELE33897100 Tel No. : ..."
# #         r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",

# #                 # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
# #                 # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
# #         r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",

# #                 # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
# #         r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

# #                 # --- General Patterns (Made more robust) ---

# #                 # General "Booking No." - now allows '/' and '-' and requires 6+ chars
# #                 # Catches Case 2 (KMTC) and others
# #         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-]{6,})",

# #                 # General "Booking Number" - FIXED BUG and made robust
# #                 # Catches Case 1 (HMM), Case 3 (ANL) and others
# #         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
# #                 # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
# #                 # E.g., "Booking Number : DELE33897100 Tel No. : ..."
# #         r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",

# #                 # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
# #                 # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
# #         r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",

# #                 # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
# #         r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

# #                 # --- General Patterns (Made more robust) ---

# #                 # General "Booking No." - now allows '/' and '-' and requires 6+ chars
# #                 # Catches Case 2 (KMTC) and others
# #         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-])",

# #                 # General "Booking Number" - FIXED BUG and made robust
# #                 # Catches Case 1 (HMM), Case 3 (ANL) and others
# #         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
# #         r"Booking\s+No\s*/\s*Ref\.?\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",

# #                 # Covers: Booking Number : DELE33897100
# #         r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",

# #                 # Covers: Booking Confirmation...Booking Number: ABP0158454
# #         r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9])",

# #                 # Covers: Booking Notice...Booking Number : DELE33897100 (again)
# #         r"Booking\s+Notice.*?Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",
# #                 # Booking No / Ref. No
# #         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
# #         r"Booking\s*Number\.?\s*[:\-]?\s*([A-Z0-9]+)",
# #         r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]+)",
# #         r"Booking\s*Acknowledgement\.?\s*([0-9]+)",
# #         r"BOOKING\s*REFERENCE\s*[:\-]?\s*([A-Z0-9]+)",
# #         r"1\*\s*Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
# #         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]+)",
# #         # Maersk (MSK) specific patterns
# #         r"Booking\s*No\.\s*:\s*([0-9]{8,})",  # "Booking No.: 250884402"
# #         r"Booking\s*No\.\s*([0-9]{8,})\s+Print\s*Date",

# #         # Hapag Lloyd specific patterns
# #         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]{6,})",
# #         r"Our\s*Ref\.\s*[:\-]?\s*([A-Z0-9]{6,})",
# #         # Portal/Reference patterns
# #         r"Portal\s*Booking\s*Ref\s*[:\-]?\s*([A-Z0-9/\-]+)",

# #         # KMTC specific patterns
# #         r"Booking\s+No\.\s+([A-Z0-9]+)\s+Booking\s+Date",
# #         r"Booking\s+No\.\s+([A-Z0-9]+)",


# #         # ANL/CMA CGM specific patterns with contextual landmarks
# #         r"Booking\s*Number\s*:\s*([A-Z0-9]+)\s+Bkg\s*Pty\s*Ref",
# #         r"Booking\s*Number\s*:\s*([A-Z0-9]+)\s+Booking\s*Date",

# #         # PIL specific patterns
# #         r"Booking\s*No\s*:\s*([A-Z0-9]+)",
# #         r"BKG\s*NO\s*:\s*([A-Z0-9]+)",

# #         # General patterns with Tel/Phone number context
# #         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]+)\s+Tel\s*No",
# #         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]+)\s+Phone",

# #         # Booking Reference variations
# #         r"Booking\s*Reference\s*No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",
# #         r"Booking\s*Ref\.\s*No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

# #         # Notice/Confirmation specific patterns
# #         r"Booking\s*Notice.*?Booking\s*No\.\s*([A-Z0-9]+)",
# #         r"Booking\s*Confirmation.*?Booking\s*Number\s*:\s*([A-Z0-9]+)",

# #         # Export reference patterns (sometimes used as booking ref)
# #         r"Export\s*Ref\.?\s*NO\s*:\s*([A-Z0-9]+)",

# #         # General robust patterns (fallback)
# #         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]{6,})",
# #         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",

# #         # Case insensitive general patterns
# #         r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
# #         r"BOOKING\s*NO\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
# #         r"BOOKING\s+NUMBER:\s+([0-9]{10})",

# #         # Reference with numbers
# #         r"Reference\s*[:\-]?\s*([A-Z0-9]{8,})",
# #         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]{6,})",

# #         # Alphanumeric patterns with minimum length
# #         r"(?:Booking|BKG)\s*(?:Number|No\.?|Ref\.?)\s*[:\-]?\s*([A-Z0-9]{6,})",

# #         # Last resort - any alphanumeric after booking keywords
# #         r"Booking.*?([A-Z]{3}[0-9]{6,})",  # Pattern like DEL500127800
# #         r"Booking.*?([A-Z0-9]{10,})",  # Long alphanumeric codes
# #     ]

# #     # Clean the text - remove extra whitespace and normalize
# #     text = re.sub(r'\s+', ' ', text)

# #     for pattern in patterns:
# #         match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
# #         if match:
# #             booking_no = match.group(1).strip()
# #             # Additional validation - ensure it's not just numbers or too short for most patterns
# #             # Exception: Allow pure numbers for Maersk booking numbers (8+ digits)
# #             if len(booking_no) >= 6:
# #                 if (re.search(r'[A-Z]', booking_no) and re.search(r'[0-9]', booking_no)) or \
# #                         (booking_no.isdigit() and len(booking_no) >= 8):
# #                     return booking_no

# #     return "Not Found"


# # # Test function with sample texts from your documents
# # def test_extraction():
# #     """Test the extraction function with sample booking confirmations"""

# #     test_cases = [
# #         # KMTC
# #         "Booking No. IN00487801 Booking Date 2025.05.29",
# #         # ANL
# #         "Booking Number: ABP0158454 Bkg Pty Ref: 04-Jun-25",
# #         # PIL
# #         "Booking No : DEL500127800",
# #         # CMA CGM
# #         "Booking Number: CAD0803097 Bkg Pty Ref: 28-Feb-25",
# #         # Maersk (MSK)
# #         "Booking No.: 250884402 Print Date:",
# #         # Hapag Lloyd (example)
# #         "Our Reference: HLCUXYZ123456",
# #     ]

# #     print("Testing booking number extraction:")
# #     for i, text in enumerate(test_cases, 1):
# #         result = extract_booking_number(text)
# #         print(f"Test {i}: {result}")

# # ------------------ Main App ------------------
# st.set_page_config(page_title="OCR PDF Extractor + DOCX Filler", layout="wide")
# st.sidebar.success(f"✅ Logged in as {st.session_state['user']}")
# if st.sidebar.button("🚪 Logout"):
#     del st.session_state["user"]
#     st.rerun()

# st.title("📄 OCR PDF Extractor & DOCX Auto-Filler")

# # Poppler path (for macOS)
# poppler_path = "/opt/homebrew/bin"
# os.environ["PATH"] += os.pathsep + poppler_path

# # ---------- Extraction Functions ----------
# def extract_ticket_fields(text):
#     data = {}
#     if match := re.search(r"TICKET NO\s*[:\-]?\s*(.*?),\s*(DICT\d+)", text):
#         data["TICKET NO"] = match.group(2)
#     if match := re.search(r"CUSTOMER\s*[:\-]?\s*(.+?)(?:\s*ADDRESS|[\n\r])", text, re.IGNORECASE):
#         data["CUSTOMER"] = match.group(1).strip()
#     if match := re.search(r"CONTAINER NO\s*[:\-]?\s*([A-Z0-9]+)", text):
#         data["CONTAINER NO"] = match.group(1)
#     if match := re.search(r"DATE OUT\s*[:\-]?\s*(\d{2}[-/]\d{2}[-/]\d{4})", text):
#         data["DATE OUT"] = match.group(1)
#     if match := re.search(r"TIME\s*OUT\s*[:\-]?\s*\(?(\d{1,2}[:]\d{2}[:]\d{2})", text, re.IGNORECASE):
#         data["TIME OUT"] = match.group(1)
#     if match := re.search(r"NET CARGO WEIGHT\s*[:\-]?\s*([0-9]+\.[0-9]+)", text):
#         data["NET CARGO WEIGHT"] = match.group(1)
#     if match := re.search(r"CONTAINER TARE WT TOTAL\s*[:\-]?\s*([0-9]+\.[0-9]+)", text):
#         data["CONTAINER TARE WT TOTAL"] = match.group(1)
#     if match := re.search(r"GROSS WEIGHT\s*[:\-]?\s*([0-9]+\.[0-9]+)", text):
#         data["GROSS WEIGHT"] = match.group(1)
#     if match := re.search(r"SIZE\s*[:\-]?\s*([0-9]+)", text):
#         data["SIZE"] = match.group(1)
#     if match := re.search(r"MAX\s*GW\s*\(CNTR\s*\)\s*([0-9]+\.[0-9]+)", text, re.IGNORECASE):
#         data["MAX GW (CNTR)"] = match.group(1)
#     return data


# def extract_booking_number(text):
#     """
#     Extract booking numbers from various shipping line booking confirmations.
#     Supports multiple formats from different carriers like KMTC, ANL, PIL, CMA CGM, etc.
#     """
#     patterns = [
#         r"Portal\s*Booking\s*Ref\s*[:\-]?\s*(.*)",

#                 # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
#                 # E.g., "Booking Number : DELE33897100 Tel No. : ..."
#         r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",

#                 # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
#                 # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
#         r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",

#                 # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
#         r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

#                 # --- General Patterns (Made more robust) ---

#                 # General "Booking No." - now allows '/' and '-' and requires 6+ chars
#                 # Catches Case 2 (KMTC) and others
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-]{6,})",

#                 # General "Booking Number" - FIXED BUG and made robust
#                 # Catches Case 1 (HMM), Case 3 (ANL) and others
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
#                 # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
#                 # E.g., "Booking Number : DELE33897100 Tel No. : ..."
#         r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",

#                 # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
#                 # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
#         r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",

#                 # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
#         r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

#                 # --- General Patterns (Made more robust) ---

#                 # General "Booking No." - now allows '/' and '-' and requires 6+ chars
#                 # Catches Case 2 (KMTC) and others
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-])",

#                 # General "Booking Number" - FIXED BUG and made robust
#                 # Catches Case 1 (HMM), Case 3 (ANL) and others
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
#         r"Booking\s+No\s*/\s*Ref\.?\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",

#                 # Covers: Booking Number : DELE33897100
#         r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",

#                 # Covers: Booking Confirmation...Booking Number: ABP0158454
#         r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9])",

#                 # Covers: Booking Notice...Booking Number : DELE33897100 (again)
#         r"Booking\s+Notice.*?Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",
#                 # Booking No / Ref. No
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"Booking\s*Number\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"Booking\s*Acknowledgement\.?\s*([0-9]+)",
#         r"BOOKING\s*REFERENCE\s*[:\-]?\s*([A-Z0-9]+)",
#         r"1\*\s*Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]+)",
#         # Maersk (MSK) specific patterns
#         r"Booking\s*No\.\s*:\s*([0-9]{8,})",  # "Booking No.: 250884402"
#         r"Booking\s*No\.\s*([0-9]{8,})\s+Print\s*Date",

#         # Hapag Lloyd specific patterns
#         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]{6,})",
#         r"Our\s*Ref\.\s*[:\-]?\s*([A-Z0-9]{6,})",
#         # Portal/Reference patterns
#         r"Portal\s*Booking\s*Ref\s*[:\-]?\s*([A-Z0-9/\-]+)",

#         # KMTC specific patterns
#         r"Booking\s+No\.\s+([A-Z0-9]+)\s+Booking\s+Date",
#         r"Booking\s+No\.\s+([A-Z0-9]+)",


#         # ANL/CMA CGM specific patterns with contextual landmarks
#         r"Booking\s*Number\s*:\s*([A-Z0-9]+)\s+Bkg\s*Pty\s*Ref",
#         r"Booking\s*Number\s*:\s*([A-Z0-9]+)\s+Booking\s*Date",

#         # PIL specific patterns
#         r"Booking\s*No\s*:\s*([A-Z0-9]+)",
#         r"BKG\s*NO\s*:\s*([A-Z0-9]+)",

#         # General patterns with Tel/Phone number context
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]+)\s+Tel\s*No",
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]+)\s+Phone",

#         # Booking Reference variations
#         r"Booking\s*Reference\s*No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",
#         r"Booking\s*Ref\.\s*No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

#         # Notice/Confirmation specific patterns
#         r"Booking\s*Notice.*?Booking\s*No\.\s*([A-Z0-9]+)",
#         r"Booking\s*Confirmation.*?Booking\s*Number\s*:\s*([A-Z0-9]+)",

#         # Export reference patterns (sometimes used as booking ref)
#         r"Export\s*Ref\.?\s*NO\s*:\s*([A-Z0-9]+)",

#         # General robust patterns (fallback)
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]{6,})",
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",

#         # Case insensitive general patterns
#         r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
#         r"BOOKING\s*NO\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
#         r"BOOKING\s+NUMBER:\s+([0-9]{10})",

#         # Reference with numbers
#         r"Reference\s*[:\-]?\s*([A-Z0-9]{8,})",
#         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]{6,})",

#         # Alphanumeric patterns with minimum length
#         r"(?:Booking|BKG)\s*(?:Number|No\.?|Ref\.?)\s*[:\-]?\s*([A-Z0-9]{6,})",

#         # Last resort - any alphanumeric after booking keywords
#         r"Booking.*?([A-Z]{3}[0-9]{6,})",  # Pattern like DEL500127800
#         r"Booking.*?([A-Z0-9]{10,})",  # Long alphanumeric codes
#     ]

#     # Clean the text - remove extra whitespace and normalize
#     text = re.sub(r'\s+', ' ', text)

#     for pattern in patterns:
#         match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
#         if match:
#             booking_no = match.group(1).strip()
#             # Additional validation - ensure it's not just numbers or too short for most patterns
#             # Exception: Allow pure numbers for Maersk booking numbers (8+ digits)
#             if len(booking_no) >= 6:
#                 if (re.search(r'[A-Z]', booking_no) and re.search(r'[0-9]', booking_no)) or \
#                         (booking_no.isdigit() and len(booking_no) >= 8):
#                     return booking_no

#     return "Not Found"


# # Test function with sample texts from your documents
# def test_extraction():
#     """Test the extraction function with sample booking confirmations"""

#     test_cases = [
#         # KMTC
#         "Booking No. IN00487801 Booking Date 2025.05.29",
#         # ANL
#         "Booking Number: ABP0158454 Bkg Pty Ref: 04-Jun-25",
#         # PIL
#         "Booking No : DEL500127800",
#         # CMA CGM
#         "Booking Number: CAD0803097 Bkg Pty Ref: 28-Feb-25",
#         # Maersk (MSK)
#         "Booking No.: 250884402 Print Date:",
#         # Hapag Lloyd (example)
#         "Our Reference: HLCUXYZ123456",
#     ]

#     print("Testing booking number extraction:")
#     for i, text in enumerate(test_cases, 1):
#         result = extract_booking_number(text)
#         print(f"Test {i}: {result}")


# #For google drive sync option 
# # TEMPLATE_DIR = r"C:\Users\Piyanshu\Google Drive\OCR_Templates"
# # TEMPLATE_DIR = "/Users/piyanshu/Desktop/shared folder/Templates"

# # def fill_docx_template(customer_name, data, booking_number, output_dir=None):
# #     if output_dir is None:
# #         desktop_path = "/Users/piyanshu/Desktop"
# #         output_dir = os.path.join(desktop_path, datetime.today().strftime("%Y-%m-%d"))

# #     customer_name_clean = customer_name.upper().strip()
# #     docx_path = os.path.join(TEMPLATE_DIR, f"{customer_name_clean}.docx")
# #     doc_path = os.path.join(TEMPLATE_DIR, f"{customer_name_clean}.doc")

# #     if os.path.exists(docx_path):
# #         template_path = docx_path
# #     elif os.path.exists(doc_path):
# #         template_path = doc_path
# #     else:
# #         return f"❌ Template not found for '{customer_name}'."
# def fill_docx_template(customer_name, data, booking_number, output_dir=None):
#     if output_dir is None:
#         desktop_path = "/Users/piyanshu/Desktop"
#         output_dir = os.path.join(desktop_path, datetime.today().strftime("%Y-%m-%d"))

#     customer_name_clean = customer_name.upper().strip()
#     docx_path = f"templates/{customer_name_clean}.docx"
#     doc_path = f"templates/{customer_name_clean}.doc"

#     if os.path.exists(docx_path):
#         template_path = docx_path
#     elif os.path.exists(doc_path):
#         template_path = doc_path
#     else:
#         return f"❌ Template not found for '{customer_name}'."

#     doc = Document(template_path)
#     date_time_weighing = f"{data.get('DATE OUT', '')}      {data.get('TIME OUT', '')}"
#     today = datetime.today().strftime("%d.%m.%Y")

#     replacements = {
#         "Booking No.": booking_number,
#         "Container No.": data.get("CONTAINER NO", ""),
#         "Container Size (TEU/FEU/other)": data.get("SIZE", ""),
#         "Maximum permissible  weight of container as per the CSC plate": data.get("MAX GW (CNTR)", ""),
#         "Weighing slip no.": data.get("TICKET NO", ""),
#         "Date and time of weighing": date_time_weighing,
#     }

#     weight_field_variants = {
#         ("CARGO WT", "CARGO WEIGHT"): data.get("NET CARGO WEIGHT", ""),
#         ("TARE WT", "TARE WEIGHT", "TARE  WT", "TARE  WEIGHT"): data.get("CONTAINER TARE WT TOTAL", ""),
#         ("VGM WT", "VGM WEIGHT", "VGM  WT", "VGM  WEIGHT"): data.get("GROSS WEIGHT", ""),
#     }

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 original_text = cell.text

#                 # Normal replacements
#                 for key, val in replacements.items():
#                     if key.lower() in original_text.lower():
#                         if len(row.cells) > 2 and key.lower() in row.cells[1].text.lower():
#                             row.cells[2].text = str(val)

#                 # Robust weight replacements
#                 for keys_tuple, weight_val in weight_field_variants.items():
#                     if any(lbl.lower() in original_text.lower() for lbl in keys_tuple):
#                         lines = cell.text.splitlines()
#                         new_lines = []
#                         for line in lines:
#                             matched_label = next((lbl for lbl in keys_tuple if lbl.lower() in line.lower()), None)
#                             if matched_label:
#                                 new_lines.append(f"{matched_label} :    {weight_val}")
#                             else:
#                                 new_lines.append(line)
#                         cell.text = "\n".join(new_lines)

#     for para in doc.paragraphs:
#         if "DT." in para.text:
#             para.text = para.text.replace("DT.", f"DT. {today}")
#             break

#     os.makedirs(output_dir, exist_ok=True)
#     timestamp = datetime.now().strftime("%H%M%S")
#     filename = f"filled_{customer_name.replace(' ', '_')}_{timestamp}.docx"
#     output_path = os.path.join(output_dir, filename)
#     doc.save(output_path)
#     return output_path

# # ---------- PDF Uploads ----------
# st.header("🎫 Ticket PDF Upload")
# ticket_pdf = st.file_uploader("Upload Ticket PDF", type="pdf", key="ticket")

# ticket_data = {}
# if ticket_pdf:
#     with st.spinner("Processing Ticket PDF..."):
#         images = convert_from_bytes(ticket_pdf.read(), dpi=300)
#     ticket_text = pytesseract.image_to_string(images[0])
#     st.image(images[0], caption="Ticket Page 1", use_container_width=True)
#     ticket_data = extract_ticket_fields(ticket_text)
#     st.subheader("🧾 Extracted Ticket Fields")
#     st.json(ticket_data)

# st.header("📘 Booking PDF Upload")
# booking_pdf = st.file_uploader("Upload Booking PDF", type="pdf", key="booking")

# booking_no = ""
# if booking_pdf:
#     with st.spinner("Processing Booking PDF..."):
#         images = convert_from_bytes(booking_pdf.read(), dpi=300)
#     booking_text = pytesseract.image_to_string(images[0])
#     st.image(images[0], caption="Booking Page 1", use_container_width=True)
#     booking_no = extract_booking_number(booking_text)
#     st.subheader("🔖 Booking Number")

#     # Make booking number editable
#     booking_no = st.text_input("Edit Booking Number if needed:", value=booking_no)


# # ---------- Generate DOCX ----------
# if ticket_pdf and booking_pdf and ticket_data.get("CUSTOMER"):
#     customer_name = ticket_data["CUSTOMER"]
#     st.header("📄 Generate Filled DOCX")

#     if st.button("Generate Document"):
#         desktop_path = "/Users/piyanshu/Desktop"
#         output_dir = os.path.join(desktop_path, datetime.today().strftime("%Y-%m-%d"))
#         result = fill_docx_template(customer_name, ticket_data, booking_no, output_dir)
#         if result.endswith(".docx"):
#             with open(result, "rb") as f:
#                 st.download_button(
#                     label="📥 Download Filled DOCX",
#                     data=f,
#                     file_name=os.path.basename(result),
#                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                 )
#             st.success(f"✅ File saved at: `{result}`")
#             st.info(f"🗂️ All files for today are saved in folder: `{output_dir}`")
#         else:
#             st.error(result)

# # ---------- ZIP Download ----------
# today_folder = datetime.today().strftime("%Y-%m-%d")
# if os.path.exists(today_folder):
#     zip_path = f"{today_folder}.zip"
#     with zipfile.ZipFile(zip_path, "w") as zipf:
#         for root, _, files in os.walk(today_folder):
#             for file in files:
#                 file_path = os.path.join(root, file)
#                 zipf.write(file_path, os.path.relpath(file_path, today_folder))

#     # with open(zip_path, "rb") as f:
#         # st.download_button("📦 Download All DOCX as ZIP", f, file_name=os.path.basename(zip_path))




import streamlit as st
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import os
import re
from datetime import datetime
from docx import Document
import zipfile


st.title("📄 OCR PDF Extractor & DOCX Auto-Filler")

# Poppler path (for macOS)
poppler_path = "/opt/homebrew/bin"
os.environ["PATH"] += os.pathsep + poppler_path

# ---------- Extraction Functions ----------
def extract_ticket_fields(text):
    data = {}
    if match := re.search(r"TICKET NO\s*[:\-]?\s*(.*?),\s*(DICT\d+)", text):
        data["TICKET NO"] = match.group(2)
    if match := re.search(r"CUSTOMER\s*[:\-]?\s*(.+?)(?:\s*ADDRESS|[\n\r])", text, re.IGNORECASE):
        data["CUSTOMER"] = match.group(1).strip()
    if match := re.search(r"CONTAINER NO\s*[:\-]?\s*([A-Z0-9]+)", text):
        data["CONTAINER NO"] = match.group(1)
    if match := re.search(r"DATE OUT\s*[:\-]?\s*(\d{2}[-/]\d{2}[-/]\d{4})", text):
        data["DATE OUT"] = match.group(1)
    if match := re.search(r"TIME\s*OUT\s*[:\-]?\s*\(?(\d{1,2}[:]\d{2}[:]\d{2})", text, re.IGNORECASE):
        data["TIME OUT"] = match.group(1)
    if match := re.search(r"NET CARGO WEIGHT\s*[:\-]?\s*([0-9]+\.[0-9]+)", text):
        data["NET CARGO WEIGHT"] = match.group(1)
    if match := re.search(r"CONTAINER TARE WT TOTAL\s*[:\-]?\s*([0-9]+\.[0-9]+)", text):
        data["CONTAINER TARE WT TOTAL"] = match.group(1)
    if match := re.search(r"GROSS WEIGHT\s*[:\-]?\s*([0-9]+\.[0-9]+)", text):
        data["GROSS WEIGHT"] = match.group(1)
    if match := re.search(r"SIZE\s*[:\-]?\s*([0-9]+)", text):
        data["SIZE"] = match.group(1)
    if match := re.search(r"MAX\s*GW\s*\(CNTR\s*\)\s*([0-9]+\.[0-9]+)", text, re.IGNORECASE):
        data["MAX GW (CNTR)"] = match.group(1)
    return data

# def extract_booking_number(text):
#     patterns = [
#          r"Portal\s*Booking\s*Ref\s*[:\-]?\s*(.*)",
#
#         # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
#         # E.g., "Booking Number : DELE33897100 Tel No. : ..."
#         r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",
#
#         # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
#         # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
#         r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",
#
#         # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
#         r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",
#
#         # --- General Patterns (Made more robust) ---
#
#         # General "Booking No." - now allows '/' and '-' and requires 6+ chars
#         # Catches Case 2 (KMTC) and others
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
#
#         # General "Booking Number" - FIXED BUG and made robust
#         # Catches Case 1 (HMM), Case 3 (ANL) and others
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
#         # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
#         # E.g., "Booking Number : DELE33897100 Tel No. : ..."
#         r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",
#
#         # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
#         # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
#         r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",
#
#         # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
#         r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",
#
#         # --- General Patterns (Made more robust) ---
#
#         # General "Booking No." - now allows '/' and '-' and requires 6+ chars
#         # Catches Case 2 (KMTC) and others
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-])",
#
#         # General "Booking Number" - FIXED BUG and made robust
#         # Catches Case 1 (HMM), Case 3 (ANL) and others
#         r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
#         r"Booking\s+No\s*/\s*Ref\.?\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
#
#         # Covers: Booking Number : DELE33897100
#         r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",
#
#         # Covers: Booking Confirmation...Booking Number: ABP0158454
#         r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9])",
#
#         # Covers: Booking Notice...Booking Number : DELE33897100 (again)
#         r"Booking\s+Notice.*?Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",
#         # Booking No / Ref. No
#         r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"Booking\s*Number\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"BOOKING\s*REFERENCE\s*[:\-]?\s*([A-Z0-9]+)",
#         r"1\*\s*Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
#         r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]+)"
#     ]
#     for pattern in patterns:
#         match = re.search(pattern, text, re.IGNORECASE)
#         if match:
#             return match.group(1)
#     return "Not Found"
import re


def extract_booking_number(text):
    """
    Extract booking numbers from various shipping line booking confirmations.
    Supports multiple formats from different carriers like KMTC, ANL, PIL, CMA CGM, etc.
    """
    patterns = [
        r"Portal\s*Booking\s*Ref\s*[:\-]?\s*(.*)",

                # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
                # E.g., "Booking Number : DELE33897100 Tel No. : ..."
        r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",

                # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
                # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
        r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",

                # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
        r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

                # --- General Patterns (Made more robust) ---

                # General "Booking No." - now allows '/' and '-' and requires 6+ chars
                # Catches Case 2 (KMTC) and others
        r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-]{6,})",

                # General "Booking Number" - FIXED BUG and made robust
                # Catches Case 1 (HMM), Case 3 (ANL) and others
        r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
                # Case 1 (HMM): Specific pattern using "Tel No." as a landmark
                # E.g., "Booking Number : DELE33897100 Tel No. : ..."
        r"Booking\s+Number\s*[:\s]+([A-Z0-9]+)\s+Tel\s+No",

                # Case 3 (ANL): Specific pattern using "Bkg Pty Ref:" as a landmark
                # E.g., "Booking Number: ABP0158454 Bkg Pty Ref:"
        r"Booking\s+Number\s*:\s*([A-Z0-9]+)\s+Bkg\s+Pty\s+Ref:",

                # Case 1 (HMM): Catches "Booking Reference No. : ..." at the bottom
        r"Booking\s+Reference\s+No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

                # --- General Patterns (Made more robust) ---

                # General "Booking No." - now allows '/' and '-' and requires 6+ chars
                # Catches Case 2 (KMTC) and others
        r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9/\-])",

                # General "Booking Number" - FIXED BUG and made robust
                # Catches Case 1 (HMM), Case 3 (ANL) and others
        r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9/\-]{6,})",
        r"Booking\s+No\s*/\s*Ref\.?\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",

                # Covers: Booking Number : DELE33897100
        r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",

                # Covers: Booking Confirmation...Booking Number: ABP0158454
        r"Booking\s+Number\s*[:\-]?\s*([A-Z0-9])",

                # Covers: Booking Notice...Booking Number : DELE33897100 (again)
        r"Booking\s+Notice.*?Booking\s+Number\s*[:\-]?\s*([A-Z0-9]{6,})",
                # Booking No / Ref. No
        r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
        r"Booking\s*Number\.?\s*[:\-]?\s*([A-Z0-9]+)",
        r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]+)",
        r"Booking\s*Acknowledgement\.?\s*([0-9]+)",
        r"BOOKING\s*REFERENCE\s*[:\-]?\s*([A-Z0-9]+)",
        r"1\*\s*Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]+)",
        r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]+)",
        # Maersk (MSK) specific patterns
        r"Booking\s*No\.\s*:\s*([0-9]{8,})",  # "Booking No.: 250884402"
        r"Booking\s*No\.\s*([0-9]{8,})\s+Print\s*Date",

        # Hapag Lloyd specific patterns
        r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]{6,})",
        r"Our\s*Ref\.\s*[:\-]?\s*([A-Z0-9]{6,})",
        # Portal/Reference patterns
        r"Portal\s*Booking\s*Ref\s*[:\-]?\s*([A-Z0-9/\-]+)",

        # KMTC specific patterns
        r"Booking\s+No\.\s+([A-Z0-9]+)\s+Booking\s+Date",
        r"Booking\s+No\.\s+([A-Z0-9]+)",


        # ANL/CMA CGM specific patterns with contextual landmarks
        r"Booking\s*Number\s*:\s*([A-Z0-9]+)\s+Bkg\s*Pty\s*Ref",
        r"Booking\s*Number\s*:\s*([A-Z0-9]+)\s+Booking\s*Date",

        # PIL specific patterns
        r"Booking\s*No\s*:\s*([A-Z0-9]+)",
        r"BKG\s*NO\s*:\s*([A-Z0-9]+)",

        # General patterns with Tel/Phone number context
        r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]+)\s+Tel\s*No",
        r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]+)\s+Phone",

        # Booking Reference variations
        r"Booking\s*Reference\s*No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",
        r"Booking\s*Ref\.\s*No\.\s*[:\-]?\s*([A-Z0-9/\-]+)",

        # Notice/Confirmation specific patterns
        r"Booking\s*Notice.*?Booking\s*No\.\s*([A-Z0-9]+)",
        r"Booking\s*Confirmation.*?Booking\s*Number\s*:\s*([A-Z0-9]+)",

        # Export reference patterns (sometimes used as booking ref)
        r"Export\s*Ref\.?\s*NO\s*:\s*([A-Z0-9]+)",

        # General robust patterns (fallback)
        r"Booking\s*Number\s*[:\-]?\s*([A-Z0-9]{6,})",
        r"Booking\s*No\.?\s*[:\-]?\s*([A-Z0-9]{6,})",

        # Case insensitive general patterns
        r"BOOKING\s*NUMBER\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
        r"BOOKING\s*NO\.?\s*[:\-]?\s*([A-Z0-9]{6,})",
        r"BOOKING\s+NUMBER:\s+([0-9]{10})",

        # Reference with numbers
        r"Reference\s*[:\-]?\s*([A-Z0-9]{8,})",
        r"Our\s*Reference\s*[:\-]?\s*([A-Z0-9]{6,})",

        # Alphanumeric patterns with minimum length
        r"(?:Booking|BKG)\s*(?:Number|No\.?|Ref\.?)\s*[:\-]?\s*([A-Z0-9]{6,})",

        # Last resort - any alphanumeric after booking keywords
        r"Booking.*?([A-Z]{3}[0-9]{6,})",  # Pattern like DEL500127800
        r"Booking.*?([A-Z0-9]{10,})",  # Long alphanumeric codes
    ]

    # Clean the text - remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', text)

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            booking_no = match.group(1).strip()
            # Additional validation - ensure it's not just numbers or too short for most patterns
            # Exception: Allow pure numbers for Maersk booking numbers (8+ digits)
            if len(booking_no) >= 6:
                if (re.search(r'[A-Z]', booking_no) and re.search(r'[0-9]', booking_no)) or \
                        (booking_no.isdigit() and len(booking_no) >= 8):
                    return booking_no

    return "Not Found"


# Test function with sample texts from your documents
def test_extraction():
    """Test the extraction function with sample booking confirmations"""

    test_cases = [
        # KMTC
        "Booking No. IN00487801 Booking Date 2025.05.29",
        # ANL
        "Booking Number: ABP0158454 Bkg Pty Ref: 04-Jun-25",
        # PIL
        "Booking No : DEL500127800",
        # CMA CGM
        "Booking Number: CAD0803097 Bkg Pty Ref: 28-Feb-25",
        # Maersk (MSK)
        "Booking No.: 250884402 Print Date:",
        # Hapag Lloyd (example)
        "Our Reference: HLCUXYZ123456",
    ]

    print("Testing booking number extraction:")
    for i, text in enumerate(test_cases, 1):
        result = extract_booking_number(text)
        print(f"Test {i}: {result}")


# Uncomment to run tests
# test_extraction()
TEMPLATE_DIR = "/Users/piyanshu/PycharmProjects/pdftoexcel/Templates"

def fill_docx_template(customer_name, data, booking_number, output_dir=None):
    if output_dir is None:
        desktop_path = "/Users/piyanshu/Desktop"
        output_dir = os.path.join(desktop_path, datetime.today().strftime("%Y-%m-%d"))

    customer_name_clean = customer_name.upper().strip()
    docx_path = os.path.join(TEMPLATE_DIR, f"{customer_name_clean}.docx")
    doc_path = os.path.join(TEMPLATE_DIR, f"{customer_name_clean}.doc")

    if os.path.exists(docx_path):
        template_path = docx_path
    elif os.path.exists(doc_path):
        template_path = doc_path
    else:
        return f"❌ Template not found for '{customer_name}'."

    doc = Document(template_path)
    date_time_weighing = f"{data.get('DATE OUT', '')}      {data.get('TIME OUT', '')}"
    today = datetime.today().strftime("%d.%m.%Y")

    replacements = {
        "Booking No.": booking_number,
        "Container No.": data.get("CONTAINER NO", ""),
        "Container Size (TEU/FEU/other)": data.get("SIZE", ""),
        "Maximum permissible  weight of container as per the CSC plate": data.get("MAX GW (CNTR)", ""),
        "Weighing slip no.": data.get("TICKET NO", ""),
        "Date and time of weighing": date_time_weighing,
    }

    weight_field_variants = {
        ("CARGO WT", "CARGO WEIGHT"): data.get("NET CARGO WEIGHT", ""),
        ("TARE WT", "TARE WEIGHT", "TARE  WT", "TARE  WEIGHT"): data.get("CONTAINER TARE WT TOTAL", ""),
        ("VGM WT", "VGM WEIGHT", "VGM  WT", "VGM  WEIGHT"): data.get("GROSS WEIGHT", ""),
    }

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text

                # Normal replacements
                for key, val in replacements.items():
                    if key.lower() in original_text.lower():
                        if len(row.cells) > 2 and key.lower() in row.cells[1].text.lower():
                            row.cells[2].text = str(val)

                # Robust weight replacements
                for keys_tuple, weight_val in weight_field_variants.items():
                    if any(lbl.lower() in original_text.lower() for lbl in keys_tuple):
                        lines = cell.text.splitlines()
                        new_lines = []
                        for line in lines:
                            matched_label = next((lbl for lbl in keys_tuple if lbl.lower() in line.lower()), None)
                            if matched_label:
                                new_lines.append(f"{matched_label} :    {weight_val}")
                            else:
                                new_lines.append(line)
                        cell.text = "\n".join(new_lines)

    for para in doc.paragraphs:
        if "DT." in para.text:
            para.text = para.text.replace("DT.", f"DT. {today}")
            break

    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%H%M%S")
    filename = f"filled_{customer_name.replace(' ', '_')}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path

# ---------- PDF Uploads ----------
st.header("🎫 Ticket PDF Upload")
ticket_pdf = st.file_uploader("Upload Ticket PDF", type="pdf", key="ticket")

ticket_data = {}
if ticket_pdf:
    with st.spinner("Processing Ticket PDF..."):
        images = convert_from_bytes(ticket_pdf.read(), dpi=300)
    ticket_text = pytesseract.image_to_string(images[0])
    st.image(images[0], caption="Ticket Page 1", use_container_width=True)
    ticket_data = extract_ticket_fields(ticket_text)
    st.subheader("🧾 Extracted Ticket Fields")
    st.json(ticket_data)

st.header("📘 Booking PDF Upload")
booking_pdf = st.file_uploader("Upload Booking PDF", type="pdf", key="booking")

booking_no = ""
if booking_pdf:
    with st.spinner("Processing Booking PDF..."):
        images = convert_from_bytes(booking_pdf.read(), dpi=300)
    booking_text = pytesseract.image_to_string(images[0])
    st.image(images[0], caption="Booking Page 1", use_container_width=True)
    booking_no = extract_booking_number(booking_text)
    st.subheader("🔖 Booking Number")

    # Make booking number editable
    booking_no = st.text_input("Edit Booking Number if needed:", value=booking_no)

# ---------- Generate DOCX ----------
if ticket_pdf and booking_pdf and ticket_data.get("CUSTOMER"):
    customer_name = ticket_data["CUSTOMER"]
    st.header("📄 Generate Filled DOCX")

    if st.button("Generate Document"):
        desktop_path = "/Users/piyanshu/Desktop"
        output_dir = os.path.join(desktop_path, datetime.today().strftime("%Y-%m-%d"))
        result = fill_docx_template(customer_name, ticket_data, booking_no, output_dir)
        if result.endswith(".docx"):
            with open(result, "rb") as f:
                st.download_button(
                    label="📥 Download Filled DOCX",
                    data=f,
                    file_name=os.path.basename(result),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            st.success(f"✅ File saved at: `{result}`")
            st.info(f"🗂️ All files for today are saved in folder: `{output_dir}`")
        else:
            st.error(result)

# ---------- ZIP Download ----------
today_folder = datetime.today().strftime("%Y-%m-%d")
if os.path.exists(today_folder):
    zip_path = f"{today_folder}.zip"
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for root, _, files in os.walk(today_folder):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, os.path.relpath(file_path, today_folder))

    # with open(zip_path, "rb") as f:
        # st.download_button("📦 Download All DOCX as ZIP", f, file_name=os.path.basename(zip_path))

